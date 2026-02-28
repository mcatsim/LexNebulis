import json
import re
import uuid
from datetime import date, datetime
from io import BytesIO
from typing import Optional

from docxtpl import DocxTemplate
from sqlalchemy import func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.documents.models import Document
from app.documents.service import get_minio_client
from app.matters.models import Matter
from app.templates.models import DocumentTemplate, GeneratedDocument, TemplateCategory
from app.templates.schemas import DocumentTemplateUpdate


async def upload_template(
    db: AsyncSession,
    file_content: bytes,
    filename: str,
    name: str,
    description: Optional[str],
    practice_area: Optional[str],
    category: TemplateCategory,
    created_by: uuid.UUID,
) -> DocumentTemplate:
    """Validate and upload a DOCX template to MinIO, then create a DB record."""
    if not filename.lower().endswith(".docx"):
        raise ValueError("Only .docx files are supported as templates")

    storage_key = f"templates/{uuid.uuid4()}/{filename}"

    # Upload to MinIO
    client = get_minio_client()
    client.put_object(
        settings.minio_bucket,
        storage_key,
        BytesIO(file_content),
        length=len(file_content),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    template = DocumentTemplate(
        name=name,
        description=description,
        practice_area=practice_area,
        category=category,
        storage_key=storage_key,
        filename=filename,
        version=1,
        created_by=created_by,
    )
    db.add(template)
    await db.flush()
    await db.refresh(template)
    return template


async def get_templates(
    db: AsyncSession,
    page: int = 1,
    page_size: int = 25,
    practice_area: Optional[str] = None,
    category: Optional[TemplateCategory] = None,
    search: Optional[str] = None,
) -> tuple[list[DocumentTemplate], int]:
    """List templates with optional filters."""
    query = select(DocumentTemplate).where(DocumentTemplate.is_active == True)  # noqa: E712
    count_query = select(func.count(DocumentTemplate.id)).where(DocumentTemplate.is_active == True)  # noqa: E712

    if practice_area:
        query = query.where(DocumentTemplate.practice_area == practice_area)
        count_query = count_query.where(DocumentTemplate.practice_area == practice_area)

    if category:
        query = query.where(DocumentTemplate.category == category)
        count_query = count_query.where(DocumentTemplate.category == category)

    if search:
        search_filter = or_(
            DocumentTemplate.name.ilike(f"%{search}%"),
            DocumentTemplate.description.ilike(f"%{search}%"),
            DocumentTemplate.filename.ilike(f"%{search}%"),
        )
        query = query.where(search_filter)
        count_query = count_query.where(search_filter)

    total = (await db.execute(count_query)).scalar_one()
    offset = (page - 1) * page_size
    result = await db.execute(query.order_by(DocumentTemplate.created_at.desc()).offset(offset).limit(page_size))
    return result.scalars().all(), total


async def get_template(db: AsyncSession, template_id: uuid.UUID) -> Optional[DocumentTemplate]:
    """Get a single template by ID."""
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    return result.scalar_one_or_none()


def get_template_variables(storage_key: str) -> list[str]:
    """Parse a DOCX template from MinIO and extract all {{ variable }} placeholders."""
    client = get_minio_client()
    response = client.get_object(settings.minio_bucket, storage_key)
    file_bytes = response.read()
    response.close()
    response.release_conn()

    # Use python-docx to read all text from the document
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(file_bytes))

    all_text = []

    # Paragraphs
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    # Headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                all_text.append(paragraph.text)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                all_text.append(paragraph.text)

    full_text = "\n".join(all_text)

    # Extract {{ variable_name }} patterns (Jinja2 style)
    pattern = r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}"
    variables = list(set(re.findall(pattern, full_text)))
    variables.sort()
    return variables


async def build_template_context(db: AsyncSession, matter_id: uuid.UUID) -> dict[str, str]:
    """Build the full context dictionary from a matter and its related entities."""
    result = await db.execute(select(Matter).where(Matter.id == matter_id))
    matter = result.scalar_one_or_none()
    if matter is None:
        raise ValueError("Matter not found")

    context: dict[str, str] = {}

    # Client fields
    client = matter.client
    if client:
        context["client_first_name"] = client.first_name or ""
        context["client_last_name"] = client.last_name or ""
        if client.first_name and client.last_name:
            context["client_full_name"] = f"{client.first_name} {client.last_name}"
        elif client.organization_name:
            context["client_full_name"] = client.organization_name
        else:
            context["client_full_name"] = client.display_name
        context["client_organization"] = client.organization_name or ""
        context["client_email"] = client.email or ""
        context["client_phone"] = client.phone or ""
        context["client_number"] = str(client.client_number) if client.client_number else ""

        # Address fields
        addr = client.address_json or {}
        context["client_address_street"] = addr.get("street", "")
        context["client_address_city"] = addr.get("city", "")
        context["client_address_state"] = addr.get("state", "")
        context["client_address_zip"] = addr.get("zip", "")

        # Build full address
        addr_parts = [
            addr.get("street", ""),
            addr.get("city", ""),
        ]
        state_zip = " ".join(filter(None, [addr.get("state", ""), addr.get("zip", "")]))
        if state_zip:
            addr_parts.append(state_zip)
        context["client_address_full"] = ", ".join(filter(None, addr_parts))

    # Matter fields
    context["matter_number"] = str(matter.matter_number) if matter.matter_number else ""
    context["matter_title"] = matter.title or ""
    context["matter_case_number"] = matter.case_number or ""
    context["matter_court"] = matter.court_name or ""
    context["matter_jurisdiction"] = matter.jurisdiction or ""
    context["matter_type"] = matter.litigation_type.value if matter.litigation_type else ""
    context["matter_status"] = matter.status.value if matter.status else ""
    context["matter_description"] = matter.description or ""

    if matter.date_opened:
        context["matter_date_opened"] = matter.date_opened.strftime("%m/%d/%Y")
    else:
        context["matter_date_opened"] = ""

    # Attorney fields
    attorney = matter.assigned_attorney
    if attorney:
        context["attorney_first_name"] = attorney.first_name or ""
        context["attorney_last_name"] = attorney.last_name or ""
        context["attorney_full_name"] = f"{attorney.first_name} {attorney.last_name}"
        context["attorney_email"] = attorney.email or ""
    else:
        context["attorney_first_name"] = ""
        context["attorney_last_name"] = ""
        context["attorney_full_name"] = ""
        context["attorney_email"] = ""

    # Firm fields
    context["firm_name"] = settings.app_name

    # Date fields
    today = date.today()
    context["current_date"] = today.strftime("%m/%d/%Y")
    context["current_date_long"] = today.strftime("%B %d, %Y")
    context["current_year"] = str(today.year)

    # Contact fields (numbered)
    if matter.contacts:
        for idx, matter_contact in enumerate(matter.contacts, start=1):
            contact = matter_contact.contact
            if contact:
                prefix = f"contact_{idx}"
                context[f"{prefix}_name"] = f"{contact.first_name} {contact.last_name}"
                context[f"{prefix}_first_name"] = contact.first_name or ""
                context[f"{prefix}_last_name"] = contact.last_name or ""
                context[f"{prefix}_role"] = matter_contact.relationship_type or ""
                context[f"{prefix}_organization"] = contact.organization or ""
                context[f"{prefix}_email"] = contact.email or ""
                context[f"{prefix}_phone"] = contact.phone or ""

    return context


async def generate_document(
    db: AsyncSession,
    template_id: uuid.UUID,
    matter_id: uuid.UUID,
    generated_by: uuid.UUID,
    custom_overrides: Optional[dict[str, str]] = None,
) -> tuple[GeneratedDocument, Document]:
    """Generate a document from a template and matter context."""
    # Fetch the template
    template = await get_template(db, template_id)
    if template is None:
        raise ValueError("Template not found")
    if not template.is_active:
        raise ValueError("Template is not active")

    # Fetch the template file from MinIO
    minio_client = get_minio_client()
    response = minio_client.get_object(settings.minio_bucket, template.storage_key)
    template_bytes = response.read()
    response.close()
    response.release_conn()

    # Build context from matter
    context = await build_template_context(db, matter_id)

    # Apply custom overrides
    if custom_overrides:
        context.update(custom_overrides)

    # Render the template using docxtpl
    doc_template = DocxTemplate(BytesIO(template_bytes))
    doc_template.render(context)

    # Save rendered document to bytes
    output_buffer = BytesIO()
    doc_template.save(output_buffer)
    output_buffer.seek(0)
    rendered_bytes = output_buffer.read()

    # Generate output filename
    base_name = template.filename.rsplit(".", 1)[0]
    output_filename = f"{base_name}_generated_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"

    # Upload generated document to MinIO under documents/{matter_id}/
    storage_key = f"{matter_id}/{uuid.uuid4()}/{output_filename}"
    minio_client.put_object(
        settings.minio_bucket,
        storage_key,
        BytesIO(rendered_bytes),
        length=len(rendered_bytes),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Create Document record
    doc_record = Document(
        matter_id=matter_id,
        uploaded_by=generated_by,
        filename=output_filename,
        storage_key=storage_key,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=len(rendered_bytes),
        version=1,
        description=f"Generated from template: {template.name}",
    )
    db.add(doc_record)
    await db.flush()
    await db.refresh(doc_record)

    # Create GeneratedDocument record
    gen_doc = GeneratedDocument(
        template_id=template_id,
        matter_id=matter_id,
        document_id=doc_record.id,
        generated_by=generated_by,
        context_json=json.dumps(context, default=str),
    )
    db.add(gen_doc)
    await db.flush()
    await db.refresh(gen_doc)

    return gen_doc, doc_record


async def update_template(
    db: AsyncSession,
    template: DocumentTemplate,
    data: DocumentTemplateUpdate,
) -> DocumentTemplate:
    """Update template metadata."""
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(template, field, value)
    await db.flush()
    await db.refresh(template)
    return template


async def delete_template(db: AsyncSession, template: DocumentTemplate) -> None:
    """Soft delete a template by setting is_active to False."""
    template.is_active = False
    await db.flush()


async def get_generated_documents(
    db: AsyncSession,
    page: int = 1,
    page_size: int = 25,
    template_id: Optional[uuid.UUID] = None,
    matter_id: Optional[uuid.UUID] = None,
) -> tuple[list[GeneratedDocument], int]:
    """List generated documents with optional filters."""
    query = select(GeneratedDocument)
    count_query = select(func.count(GeneratedDocument.id))

    if template_id:
        query = query.where(GeneratedDocument.template_id == template_id)
        count_query = count_query.where(GeneratedDocument.template_id == template_id)

    if matter_id:
        query = query.where(GeneratedDocument.matter_id == matter_id)
        count_query = count_query.where(GeneratedDocument.matter_id == matter_id)

    total = (await db.execute(count_query)).scalar_one()
    offset = (page - 1) * page_size
    result = await db.execute(query.order_by(GeneratedDocument.created_at.desc()).offset(offset).limit(page_size))
    return result.scalars().all(), total
